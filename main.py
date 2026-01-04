import io
import os
import sys
from dataclasses import dataclass
from datetime import datetime
from typing import List, Literal
from typing import Optional
from zoneinfo import ZoneInfo

from docx import Document
from pypdf import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


TZ_NAME = "Europe/Minsk"

TEMPLATE_PDF = "template.pdf"
FONT_TTF = "Montserrat-Regular.ttf"

FONT_NAME = "Montserrat"
FONT_SIZE = 10
TEXT_COLOR_RGB = (220, 0, 0)

SALUTATION_POS = {
    "x": 75.0,
    "y": 480.0,
    "rotate_deg": 90.0,
    "align": "left",
}

GREETING_POS = {
    "x": 118.0,
    "y": 500.0,
    "rotate_deg": 90.0,
    "align": "left",
}


DOCX_TABLE_INDEX = 0
HAS_HEADER_ROW = True
SKIP_EMPTY_ROWS = True

SALUTATION_FORMAT = "{dear} {name}!"
GREETING_FORMAT = "{greeting}"


@dataclass(frozen=True)
class RowData:
    name: str
    greeting: str
    dear: str


def _norm(s: str) -> str:
    return " ".join((s or "").replace("\xa0", " ").strip().split())

from lxml import etree

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}

def _tcs_from_row_xml(tr_xml: str):
    """
    Возвращает список XML-элементов w:tc для строки.
    Учитывает кейс, когда ячейки находятся внутри w:sdt/w:sdtContent (content controls).
    Порядок сохраняется как в документе.
    """
    tr = etree.fromstring(tr_xml.encode("utf-8"))
    out = []

    for child in tr:
        if child.tag == f"{{{W_NS}}}tc":
            out.append(child)
        elif child.tag == f"{{{W_NS}}}sdt":
            tcs = child.xpath("./w:sdtContent/w:tc", namespaces=NS)
            out.extend(tcs)

    return out

def _text_from_tc(tc) -> str:
    """Достаёт текст из XML-ячейки w:tc (включая вложенные элементы)."""
    parts = tc.xpath(".//w:t/text()", namespaces=NS)
    return "".join(parts)


def find_single_docx_in_cwd_strict() -> str:
    """Строго: если docx не указан, должен быть ровно один .docx в корне."""
    docx_files = [f for f in os.listdir(".") if f.lower().endswith(".docx") and os.path.isfile(f)]
    if len(docx_files) == 1:
        return docx_files[0]
    if len(docx_files) == 0:
        raise FileNotFoundError(
            "В корне проекта не найден .docx.\n"
            "Положите Word-файл рядом с main.py или запустите: python3 main.py <файл.docx>"
        )
    raise ValueError(
        "В корне найдено несколько .docx, это запрещено в строгом режиме:\n"
        "  - " + "\n  - ".join(sorted(docx_files)) + "\n"
        "Укажите нужный явно: python3 main.py <файл.docx>"
    )


def read_rows_from_docx(docx_path: str) -> List[RowData]:
    doc = Document(docx_path)
    if not doc.tables:
        raise ValueError("В .docx не найдено ни одной таблицы. Нужна таблица с 3 колонками.")

    if DOCX_TABLE_INDEX < 0 or DOCX_TABLE_INDEX >= len(doc.tables):
        raise ValueError(f"Некорректный DOCX_TABLE_INDEX={DOCX_TABLE_INDEX}. Таблиц: {len(doc.tables)}")

    table = doc.tables[DOCX_TABLE_INDEX]
    start_idx = 1 if HAS_HEADER_ROW else 0

    out: List[RowData] = []
    for i in range(start_idx, len(table.rows)):
        tr_xml = table.rows[i]._tr.xml
        tcs = _tcs_from_row_xml(tr_xml)

        if len(tcs) < 3:
            raise ValueError(f"Строка {i + 1} таблицы содержит меньше 3 ячеек (с учетом SDT).")

        name = _norm(_text_from_tc(tcs[0]))
        greeting = _norm(_text_from_tc(tcs[1]))
        dear = _norm(_text_from_tc(tcs[2]))

        if SKIP_EMPTY_ROWS and (not name and not greeting and not dear):
            continue

        if not name:
            raise ValueError(f"Пустое 'Имя отчество' в строке {i + 1}.")
        if not greeting:
            raise ValueError(f"Пустая колонка 'С Днем рождения/С Юбилеем' в строке {i + 1}.")
        if not dear:
            raise ValueError(f"Пустая колонка 'Уважаемый/Уважаемая' в строке {i + 1}.")

        out.append(RowData(name=name, greeting=greeting, dear=dear))

    if not out:
        raise ValueError("Не найдено ни одной заполненной строки для генерации страниц.")
    return out


def output_filename_from_docx(docx_path: str) -> str:
    """ГГГГ_ММ_ДД_<ИмяWord>_Вкладыши.pdf"""
    now = datetime.now(ZoneInfo(TZ_NAME))
    stem = os.path.splitext(os.path.basename(docx_path))[0].strip()

    # минимально "безопасное" имя файла (сохраняем кириллицу)
    allowed_extra = set(" _-")
    safe_stem_chars = []
    for ch in stem:
        if ch.isalnum() or ch in allowed_extra:
            safe_stem_chars.append(ch)
    safe_stem = "".join(safe_stem_chars).replace("  ", " ").strip()

    if not safe_stem:
        safe_stem = "Word"

    return f"{now.year:04d}_{now.month:02d}_{now.day:02d}_Вкладыши.pdf"


Align = Literal["left", "center"]


def draw_text_block(
    c: canvas.Canvas,
    text: str,
    x: float,
    y: float,
    rotate_deg: float,
    align: Align,
) -> None:
    """
    Рисует один текстовый блок в точке (x,y), с поворотом rotate_deg.
    Поворот выполняется вокруг (x,y).
    """
    c.saveState()
    c.translate(float(x), float(y))
    if rotate_deg:
        c.rotate(float(rotate_deg))

    if align == "center":
        c.drawCentredString(0.0, 0.0, text)
    else:
        c.drawString(0.0, 0.0, text)

    c.restoreState()


def make_overlay_pdf(
    page_width: float,
    page_height: float,
    salutation_text: str,
    greeting_text: str,
) -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=(page_width, page_height))

    r, g, b = TEXT_COLOR_RGB
    c.setFillColorRGB(r / 255.0, g / 255.0, b / 255.0)
    c.setFont(FONT_NAME, FONT_SIZE)

    draw_text_block(
        c=c,
        text=salutation_text,
        x=float(SALUTATION_POS["x"]),
        y=float(SALUTATION_POS["y"]),
        rotate_deg=float(SALUTATION_POS["rotate_deg"]),
        align=SALUTATION_POS["align"],
    )

    # Надпись 2: "С Днем рождения / С Юбилеем"
    draw_text_block(
        c=c,
        text=greeting_text,
        x=float(GREETING_POS["x"]),
        y=float(GREETING_POS["y"]),
        rotate_deg=float(GREETING_POS["rotate_deg"]),
        align=GREETING_POS["align"],
    )

    c.showPage()
    c.save()
    return buf.getvalue()


def generate_pdf(docx_path: str) -> str:
    if not os.path.exists(TEMPLATE_PDF):
        raise FileNotFoundError(f"Не найден {TEMPLATE_PDF} в корне проекта.")
    if not os.path.exists(FONT_TTF):
        raise FileNotFoundError(f"Не найден {FONT_TTF} в корне проекта.")

    # 1) Читаем Word
    rows = read_rows_from_docx(docx_path)

    # 2) Регистрируем шрифт
    pdfmetrics.registerFont(TTFont(FONT_NAME, FONT_TTF))

    # 3) Читаем template.pdf в байты, чтобы на каждую строку брать "свежую" страницу 2
    with open(TEMPLATE_PDF, "rb") as tf:
        template_bytes = tf.read()

    reader0 = PdfReader(io.BytesIO(template_bytes))
    if len(reader0.pages) < 2:
        raise ValueError("template.pdf должен содержать минимум 2 страницы (1-я статичная, 2-я — образец).")


    page1 = reader0.pages[0]

    # Размеры берём со второй страницы
    base0 = reader0.pages[1]
    page_width = float(base0.mediabox.width)
    page_height = float(base0.mediabox.height)

    # 4) Пишем результат
    writer = PdfWriter()
    writer.add_page(page1)

    for row in rows:
        salutation_text = SALUTATION_FORMAT.format(dear=row.dear, name=row.name, greeting=row.greeting)
        greeting_text = GREETING_FORMAT.format(dear=row.dear, name=row.name, greeting=row.greeting)

        overlay_bytes = make_overlay_pdf(
            page_width=page_width,
            page_height=page_height,
            salutation_text=salutation_text,
            greeting_text=greeting_text,
        )
        overlay_reader = PdfReader(io.BytesIO(overlay_bytes))
        overlay_page = overlay_reader.pages[0]

        fresh_reader = PdfReader(io.BytesIO(template_bytes))
        fresh_base_page = fresh_reader.pages[1]

        new_page = writer.add_page(fresh_base_page)
        new_page.merge_page(overlay_page)

    out_path = output_filename_from_docx(docx_path)
    with open(out_path, "wb") as f:
        writer.write(f)

    return out_path


def main() -> None:
    if len(sys.argv) >= 2:
        docx_path = sys.argv[1]
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"Не найден Word-файл: {docx_path}")
    else:
        docx_path = find_single_docx_in_cwd_strict()

    out = generate_pdf(docx_path)
    print(f"Готово: {out}")


if __name__ == "__main__":
    main()